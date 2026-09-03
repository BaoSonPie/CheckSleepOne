from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from docx import Document

import os
import time
import threading

# ==========================================
# CẤU HÌNH
# ==========================================

WATCH_FOLDER = os.path.dirname(os.path.abspath(__file__))

OUTPUT_DOCX = os.path.join(WATCH_FOLDER, "python_code_snapshot_EAR.docx")


# ==========================================
# XỬ LÝ FILE PYTHON
# ==========================================


class PyFileHandler(FileSystemEventHandler):

    def __init__(self):
        self.timer = None

    def export_to_word(self):

        doc = Document()

        doc.add_heading("CheckSleepOne - Python Source Code", level=1)

        doc.add_paragraph(
            f"Thời gian cập nhật: " f"{time.strftime('%Y-%m-%d %H:%M:%S')}"
        )

        # Duyệt toàn bộ project
        for root, dirs, files in os.walk(WATCH_FOLDER):

            # Bỏ qua các thư mục không cần thiết
            dirs[:] = [
                d
                for d in dirs
                if d not in [".venv", "venv", "__pycache__", ".git", "runs", "dataset"]
            ]

            for file in sorted(files):

                if not file.endswith(".py"):
                    continue

                # Không đưa chính file export này vào Word
                if file == "export_py_to_word_realtime.py":
                    continue

                path = os.path.join(root, file)

                # Tên đường dẫn tương đối
                relative_path = os.path.relpath(path, WATCH_FOLDER)

                doc.add_heading(relative_path, level=2)

                try:

                    with open(path, "r", encoding="utf-8") as f:

                        code = f.read()

                    # Dùng style code
                    paragraph = doc.add_paragraph()

                    run = paragraph.add_run(code)

                    run.font.name = "Consolas"
                    run.font.size = None

                except Exception as e:

                    doc.add_paragraph(f"Lỗi đọc file: {e}")

        doc.save(OUTPUT_DOCX)

        print(f"[{time.strftime('%H:%M:%S')}] " f"Đã cập nhật Word")

    # ======================================
    # CHỜ 3 GIÂY SAU LẦN SỬA CUỐI
    # ======================================

    def schedule_export(self):

        if self.timer:
            self.timer.cancel()

        self.timer = threading.Timer(3.0, self.export_to_word)

        self.timer.start()

    # ======================================
    # KHI FILE PY ĐƯỢC SỬA
    # ======================================

    def on_modified(self, event):

        if not event.is_directory and event.src_path.endswith(".py"):

            print(f"Phát hiện thay đổi: " f"{event.src_path}")

            self.schedule_export()

    # ======================================
    # KHI TẠO FILE PY MỚI
    # ======================================

    def on_created(self, event):

        if not event.is_directory and event.src_path.endswith(".py"):

            print(f"Phát hiện file mới: " f"{event.src_path}")

            self.schedule_export()


# ==========================================
# KHỞI ĐỘNG
# ==========================================

handler = PyFileHandler()

# Xuất Word ngay lần đầu
handler.export_to_word()


observer = Observer()

observer.schedule(handler, WATCH_FOLDER, recursive=True)

observer.start()


print()
print("==========================================")
print(" CHECKSLEEPONE CODE EXPORTER")
print("==========================================")
print()
print("Đang theo dõi toàn bộ file Python...")
print()
print("Mỗi khi sửa code:")
print("1. Lưu file .py")
print("2. Chờ 3 giây")
print("3. Word tự động cập nhật")
print()
print("Nhấn Ctrl + C để dừng.")
print()


try:

    while True:
        time.sleep(1)

except KeyboardInterrupt:

    print()
    print("Đang dừng...")

    observer.stop()


observer.join()

print("Đã dừng exporter.")
